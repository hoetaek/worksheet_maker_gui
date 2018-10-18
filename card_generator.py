from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Cm, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import json

class WordCardWorksheet:
    def __init__(self, card_image, word_width, path):
        self.card_image = [os.path.join(card_image, i) for i in os.listdir(card_image)]
        self.word_width = word_width
        self.path = path

    def make_worksheet(self):
        print('학습지 만드는 중입니다.')
        # write to docx file
        document = Document()
        #changing the page margins
        sections = document.sections
        left_margin = 0.3
        right_margin = 0.3
        if os.path.exists('hwp_margin_settings.json'):
            with open('hwp_margin_settings.json') as f:
                data = json.load(f)
                left_margin = data['left_margin']
                right_margin = data['right_margin']
        for section in sections:
            section.top_margin = Cm(0.3)
            section.bottom_margin = Cm(0.3)
            section.left_margin = Cm(left_margin)
            section.right_margin = Cm(right_margin)

        word_num = len(self.card_image)
        size = self.word_width
        hint_table = document.add_table(rows = (word_num+size-1)//size, cols = size, style = 'Table Grid')
        hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(hint_table.rows):
            #######################세로 길이 정하기!

            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '1000')
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                index = i*size + j
                #단어 수 만큼 반복하기
                if index < word_num:
                    for paragraph in cell.paragraphs:
                        try:
                            run = paragraph.add_run()
                            run.add_picture(self.card_image[index], width=cell.width *95/100,
                                                height=cell.width)

                        except Exception:
                            return 'error'

                        #####가운데 정렬!!
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)

        document.save(self.path+'\도블 카드 모음.hwp')


if __name__=='__main__':
    word_image = []
    dir_name = r'C:\Users\hoetaekpro\Desktop\dobble_6cards_picture'
    for i, j in enumerate(os.listdir(dir_name)):
        word_image.append([str(i), os.path.abspath(dir_name + '\\' + j)])
    WordCardWorksheet(word_image, 3, r'C:\Users\hoetaekpro\Desktop').make_worksheet()
