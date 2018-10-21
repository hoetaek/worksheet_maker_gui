from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Cm, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import json
from syllable import get_syllable_divided

class WordCardWorksheet:
    def __init__(self, word_image, word_width, path, syllable=False):
        self.word_image = word_image
        self.words = [word[0] for word in word_image]
        self.word_width = word_width
        self.path = path
        self.syllable = syllable

    def make_worksheet(self):
        print('학습지 만드는 중입니다.')
        # write to docx file
        document = Document()
        #changing the page margins
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(0.8)
            section.left_margin = Cm(1.3)
            section.right_margin = Cm(1.3)
        head = document.add_heading('Words Worksheet', 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists('hwp_settings.json'):
            with open('hwp_settings.json') as f:
                data = json.load(f)
                para_belong = document.add_paragraph('{}학년 {}반 이름: _______'.format(data['grade'], data['class']))
        else:
            para_belong = document.add_paragraph('__학년 __반 이름: _______')
        para_belong.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        word_num = len(self.words)
        size = self.word_width
        hint_table = document.add_table(rows = (word_num+size-1)//size * 2, cols = size, style = 'Table Grid')
        hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row in enumerate(hint_table.rows):
            #######################세로 길이 정하기!
            if i%2 == 0:
                # accessing row xml and setting tr height
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), '1000')
                trHeight.set(qn('w:hRule'), "atLeast")
                trPr.append(trHeight)
            elif i%2 == 1:
                # accessing row xml and setting tr height
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), '60')
                trHeight.set(qn('w:hRule'), "atLeast")
                trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                index = i//2*size + j
                #단어 수 만큼 반복하기
                if index < word_num:
                    if i % 2 == 1:
                        # TODO font size : 15, font bold
                        cell.text = self.words[index]
                        if self.syllable:
                            cell.text = get_syllable_divided(self.words[index])
                    for paragraph in cell.paragraphs:
                        if i % 2 == 0:
                            if self.word_image[index][1] == "None":
                                cell.text = "사진 없음"
                            else:
                                try:
                                    run = paragraph.add_run()
                                    run.add_picture(self.word_image[index][1], width=cell.width *95/100,
                                                        height=cell.width)

                                except Exception as e:
                                    print("에러 발생", e)

                        #####가운데 정렬!!
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)

        document.save(self.path+'\새로운 단어 학습지.hwp')


if __name__=='__main__':
    word_image = []
    dir_name = r'C:\Users\hoetaekpro\Desktop\dobble_6cards_picture'
    for i, j in enumerate(os.listdir(dir_name)):
        word_image.append([str(i), os.path.abspath(dir_name + '\\' + j)])
    WordCardWorksheet(word_image, 3, r'C:\Users\hoetaekpro\Desktop').make_worksheet()