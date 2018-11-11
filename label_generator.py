from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class Picture2Table:
    def __init__(self, template, words=[], korean=[], pics=[]):
        self.template = Document(template)
        self.words = words
        self.korean = korean
        self.pics = pics
        self.tabel = self.template.tables[0]
        self.get_cell_size()

    def get_cell_size(self):
        self.tabel = self.template.tables[0]
        for i, row in enumerate(self.tabel.rows):
            self.height = len(self.tabel.rows)
            self.pic_height = row.height
            for j, cell in enumerate(row.cells):
                self.width = len(row.cells)
                self.pic_width = cell.width
                break
            break
        return self.pic_width, self.pic_height, self.width, self.height


    def put_pic_to_cell(self):
        for i, row in enumerate(self.tabel.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    index = i * self.width + j
                    if index < len(self.pics):
                        run = paragraph.add_run()
                        run.add_picture(self.pics[index], width=self.pic_width // 100 * 98, height=self.pic_height // 100 * 90)
                    #####상하 방향에서 가운데 정렬
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)

    def put_word_to_cell(self):
        for i, row in enumerate(self.tabel.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    index = i * self.width + j
                    if index < len(self.pics):
                        paragraph.add_run(self.words[index])
                    #####상하 방향에서 가운데 정렬
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)

    def put_korean_to_cell(self):
        for i, row in enumerate(self.tabel.rows):
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    index = i * self.width + j
                    if index < len(self.pics):
                        paragraph.add_run(self.korean[index])
                    #####상하 방향에서 가운데 정렬
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)

if __name__=='__main__':
    import os
    pics = [os.path.join('flicker_example', f) for f in os.listdir('flicker_example')]
    pic2table = Picture2Table('template_24.docx', pics=pics)
    pic2table.put_pic_to_cell()
