from base64 import b64encode
from io import BytesIO
from xml.etree import ElementTree
from zipfile import ZipFile

import httpx
import pytest
from docx import Document
from docx.shared import Cm
from fastapi.testclient import TestClient
from PIL import Image

import backend.image_search as image_search
from backend import generators
from backend.image_search import (
    IMAGE_REQUEST_HEADERS,
    OPENVERSE_API_URL,
    build_image_search_queries,
    build_search_queries,
    normalize_openverse_images,
    provider_order,
    search_images,
    search_images_with_query,
)
from backend.main import app
from backend.schemas import ImageCandidate

client = TestClient(app)
PPTX_NS = {"p": "http://schemas.openxmlformats.org/presentationml/2006/main"}
REL_NS = {"pr": "http://schemas.openxmlformats.org/package/2006/relationships"}
R_NS_URI = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
P14_NS_URI = "http://schemas.microsoft.com/office/powerpoint/2010/main"
AUDIO_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/audio"
IMAGE_MEDIA_EXTENSIONS = (".bmp", ".gif", ".jpeg", ".jpg", ".png", ".tif", ".tiff")


def test_health_endpoint() -> None:
    response = client.get("/api/health")

    assert response.status_code == 200
    assert response.json() == {"status": "ok"}


def test_flicker_endpoint_returns_pptx() -> None:
    response = client.post(
        "/api/materials/flicker.pptx",
        json={
            "items": [{"word": "cat"}, {"word": "dog"}],
            "templates": ["word", "blank"],
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    )
    assert response.content[:2] == b"PK"
    with ZipFile(BytesIO(response.content)) as archive:
        assert "[Content_Types].xml" in archive.namelist()
        assert "ppt/presentation.xml" in archive.namelist()


def test_flicker_endpoint_embeds_item_images() -> None:
    response = client.post(
        "/api/materials/flicker.pptx",
        json={
            "items": [{"word": "cat", "image": tiny_png_data_uri(width=24, height=96)}],
            "templates": ["image"],
        },
    )

    assert response.status_code == 200
    assert archive_has_media(response.content, "ppt/media/")
    assert media_dimensions(response.content, "ppt/media/") == {(800, 600)}


def test_flicker_endpoint_adds_auto_advance_fade_transitions() -> None:
    response = client.post(
        "/api/materials/flicker.pptx",
        json={
            "items": [{"word": "cat"}],
            "templates": ["word", "blank"],
        },
    )

    assert response.status_code == 200
    slide_xmls = pptx_slide_xmls(response.content)
    assert len(slide_xmls) == 2
    for slide_xml in slide_xmls:
        slide = ElementTree.fromstring(slide_xml)
        transition = slide.find("p:transition", PPTX_NS)
        assert transition is not None
        assert transition.attrib["spd"] == "med"
        assert transition.attrib[f"{{{P14_NS_URI}}}dur"] == "700"
        assert transition.attrib["advClick"] == "0"
        assert transition.attrib["advTm"] == "1000"
        assert transition.find("p:fade", PPTX_NS) is not None


def test_flicker_endpoint_embeds_transition_sound() -> None:
    response = client.post(
        "/api/materials/flicker.pptx",
        json={
            "items": [{"word": "cat"}],
            "templates": ["word"],
        },
    )

    assert response.status_code == 200
    with ZipFile(BytesIO(response.content)) as archive:
        media_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("ppt/media/") and name.endswith(".wav")
        )
        assert media_names == ["ppt/media/audio1.wav"]
        assert archive.read(media_names[0]).startswith(b"RIFF")
        slide_xml = archive.read("ppt/slides/slide1.xml").decode("utf-8")
        slide_rel_xml = archive.read("ppt/slides/_rels/slide1.xml.rels").decode("utf-8")

    slide = ElementTree.fromstring(slide_xml)
    sound = slide.find("p:transition/p:sndAc/p:stSnd/p:snd", PPTX_NS)
    assert sound is not None
    assert sound.attrib["name"] == "whoosh.wav"

    sound_rel_id = sound.attrib[f"{{{R_NS_URI}}}embed"]
    slide_rels = ElementTree.fromstring(slide_rel_xml)
    sound_rel = slide_rels.find(f"pr:Relationship[@Id='{sound_rel_id}']", REL_NS)
    assert sound_rel is not None
    assert sound_rel.attrib["Type"] == AUDIO_REL_TYPE
    assert sound_rel.attrib["Target"] == "../media/audio1.wav"


def test_flicker_word_image_slide_adds_fade_reveal_timing() -> None:
    response = client.post(
        "/api/materials/flicker.pptx",
        json={
            "items": [{"word": "cat", "image": tiny_png_data_uri()}],
            "templates": ["word-image"],
        },
    )

    assert response.status_code == 200
    slide_xml = pptx_slide_xmls(response.content)[0]
    slide = ElementTree.fromstring(slide_xml)
    timing = slide.find("p:timing", PPTX_NS)
    assert timing is not None
    effects = {
        (effect.attrib["transition"], effect.attrib["filter"])
        for effect in timing.findall(".//p:animEffect", PPTX_NS)
    }
    assert ("out", "fade") in effects
    assert ("in", "fade") in effects
    assert any(
        time_node.attrib.get("dur") == "500" for time_node in timing.findall(".//p:cTn", PPTX_NS)
    )


def test_worksheet_endpoint_returns_docx() -> None:
    response = client.post(
        "/api/materials/worksheet.docx",
        json={
            "items": [{"word": "토끼"}, {"word": "거북이"}],
            "columns": 2,
            "grade": 3,
            "class_number": 1,
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.content[:2] == b"PK"
    assert "_______" not in docx_text(response.content)
    assert "학년" in docx_text(response.content)
    assert "반" in docx_text(response.content)
    assert "이름" in docx_text(response.content)


def test_worksheet_endpoint_embeds_item_images() -> None:
    response = client.post(
        "/api/materials/worksheet.docx",
        json={
            "items": [{"word": "토끼", "image": tiny_png_data_uri(width=96, height=24)}],
            "columns": 1,
            "grade": 3,
            "class_number": 1,
        },
    )

    assert response.status_code == 200
    assert archive_has_media(response.content, "word/media/")
    assert media_dimensions(response.content, "word/media/") == {(800, 600)}


def test_word_search_endpoint_returns_premium_docx_without_underscore_name_field() -> None:
    response = client.post(
        "/api/materials/word-search.docx",
        json={
            "words": ["토끼", "사자"],
            "grid": [["토", "끼"], ["사", "자"]],
            "hints": [{"word": "토끼"}, {"word": "사자"}],
            "grade": 3,
            "class_number": 1,
            "title": "낱말 찾기",
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    text = docx_text(response.content)
    assert "_______" not in text
    assert "학년" in text
    assert "반" in text
    assert "이름" in text


def test_word_search_endpoint_embeds_hint_images() -> None:
    response = client.post(
        "/api/materials/word-search.docx",
        json={
            "words": ["토끼"],
            "grid": [["토", "끼"]],
            "hints": [{"word": "토끼", "image": tiny_png_data_uri(width=24, height=96)}],
            "grade": 3,
            "class_number": 1,
            "title": "낱말 찾기",
        },
    )

    assert response.status_code == 200
    assert archive_has_media(response.content, "word/media/")
    assert media_dimensions(response.content, "word/media/") == {(800, 600)}
    assert docx_picture_widths(response.content) == {Cm(3.2)}


def test_dobble_endpoint_embeds_card_images() -> None:
    response = client.post(
        "/api/materials/dobble.pptx",
        json={
            "cards": [
                [
                    {"word": "cat", "image": tiny_png_data_uri(width=96, height=24)},
                    {"word": "dog"},
                    {"word": "pig"},
                ]
            ],
            "pictures_per_card": 3,
        },
    )

    assert response.status_code == 200
    assert archive_has_media(response.content, "ppt/media/")
    assert media_dimensions(response.content, "ppt/media/") == {(800, 600)}


def test_dobble_endpoint_draws_circular_game_card() -> None:
    response = client.post(
        "/api/materials/dobble.pptx",
        json={
            "cards": [
                [
                    {"word": "cat"},
                    {"word": "dog"},
                    {"word": "pig"},
                    {"word": "fox"},
                    {"word": "cow"},
                ]
            ],
            "pictures_per_card": 5,
        },
    )

    assert response.status_code == 200
    slide_xml = pptx_slide_xmls(response.content)[0]
    assert 'prst="ellipse"' in slide_xml
    assert "도블 카드 1" in slide_xml


def test_dobble_endpoint_respects_image_only_display_mode() -> None:
    response = client.post(
        "/api/materials/dobble.pptx",
        json={
            "cards": [
                [
                    {"word": "cat", "image": tiny_png_data_uri()},
                    {"word": "dog"},
                    {"word": "pig"},
                ]
            ],
            "pictures_per_card": 3,
            "display_mode": "image",
        },
    )

    assert response.status_code == 200
    slide_xml = pptx_slide_xmls(response.content)[0]
    assert "cat" not in slide_xml
    assert archive_has_media(response.content, "ppt/media/")


def test_dobble_endpoint_uses_initial_fallback_for_missing_image_only_symbols() -> None:
    response = client.post(
        "/api/materials/dobble.pptx",
        json={
            "cards": [[{"word": "dog"}]],
            "pictures_per_card": 3,
            "display_mode": "image",
        },
    )

    assert response.status_code == 200
    slide_xml = pptx_slide_xmls(response.content)[0]
    assert "<a:t>D</a:t>" in slide_xml
    assert "dog" not in slide_xml


def test_dobble_endpoint_rotates_symbols_for_multi_sided_cards() -> None:
    response = client.post(
        "/api/materials/dobble.pptx",
        json={
            "cards": [
                [
                    {"word": "cat"},
                    {"word": "dog"},
                    {"word": "pig"},
                    {"word": "fox"},
                ]
            ],
            "pictures_per_card": 4,
            "display_mode": "image",
        },
    )

    assert response.status_code == 200
    slide_xml = pptx_slide_xmls(response.content)[0]
    assert 'cx="2121408"' in slide_xml
    assert 'cy="2121408"' in slide_xml
    assert 'rot="2880000"' in slide_xml
    assert 'rot="7680000"' in slide_xml
    assert 'rot="13260000"' in slide_xml
    assert 'rot="18540000"' in slide_xml


def test_image_search_validates_query() -> None:
    response = client.get("/api/images/search", params={"query": "", "limit": 3})

    assert response.status_code == 422


def test_image_search_uses_current_public_hosts_and_request_headers() -> None:
    assert OPENVERSE_API_URL == "https://api.openverse.org/v1/images/"
    assert "worksheet-maker-gui" in IMAGE_REQUEST_HEADERS["User-Agent"]
    assert IMAGE_REQUEST_HEADERS["Api-User-Agent"].startswith("worksheet-maker-gui/")


def test_search_query_variants_remove_file_or_media_suffixes() -> None:
    assert build_search_queries("토끼 png") == ["토끼 png", "토끼"]
    assert build_search_queries("rabbit photo") == ["rabbit photo", "rabbit"]
    assert build_search_queries("  cat  ") == ["cat"]


@pytest.mark.anyio
async def test_image_search_queries_use_external_translation_without_local_dictionary(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async def fake_translate(query: str) -> str | None:
        return {"거북이": "turtle", "토끼": "rabbit"}.get(query)

    monkeypatch.setattr(image_search, "translate_korean_query_to_english", fake_translate)

    assert await build_image_search_queries("거북이") == ["turtle", "거북이"]
    assert await build_image_search_queries("토끼 png") == ["rabbit", "토끼 png", "토끼"]


def test_auto_provider_prefers_commons_for_korean_queries() -> None:
    assert provider_order("토끼", "auto") == ["commons", "openverse"]
    assert provider_order("rabbit", "auto") == ["openverse", "commons"]
    assert provider_order("토끼", "openverse") == ["openverse", "commons"]


@pytest.mark.anyio
async def test_search_images_retries_clean_query_then_secondary_provider(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    openverse_queries: list[str] = []
    commons_queries: list[str] = []

    async def fake_openverse(query: str, limit: int = 6) -> list[ImageCandidate]:
        openverse_queries.append(query)
        return []

    async def fake_commons(query: str, limit: int = 6) -> list[ImageCandidate]:
        commons_queries.append(query)
        if query == "토끼":
            return [
                ImageCandidate(
                    id="commons:rabbit",
                    title="Rabbit",
                    image_url="https://example.com/rabbit.jpg",
                    thumbnail_url="https://example.com/rabbit-thumb.jpg",
                    source_url="https://commons.wikimedia.org/wiki/File:Rabbit.jpg",
                    provider="commons",
                )
            ]

        return []

    async def fake_translate(query: str) -> str | None:
        return "rabbit" if query == "토끼" else None

    monkeypatch.setattr(image_search, "search_openverse_images", fake_openverse)
    monkeypatch.setattr(image_search, "search_commons_images", fake_commons)
    monkeypatch.setattr(image_search, "translate_korean_query_to_english", fake_translate)

    results = await search_images("토끼 png", 3, "openverse")

    assert [candidate.provider for candidate in results] == ["commons"]
    assert openverse_queries == ["rabbit", "토끼 png", "토끼"]
    assert commons_queries == ["rabbit", "토끼 png", "토끼"]


@pytest.mark.anyio
async def test_search_images_reports_the_query_that_produced_results(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    async def fake_openverse(query: str, limit: int = 6) -> list[ImageCandidate]:
        if query == "turtle":
            return [
                ImageCandidate(
                    id="openverse:turtle",
                    title="Turtle",
                    image_url="https://example.com/turtle.jpg",
                    thumbnail_url="https://example.com/turtle-thumb.jpg",
                    source_url="https://openverse.org/image/turtle",
                    provider="openverse",
                )
            ]

        return []

    async def fake_commons(query: str, limit: int = 6) -> list[ImageCandidate]:
        return []

    async def fake_translate(query: str) -> str | None:
        return "turtle" if query == "거북이" else None

    monkeypatch.setattr(image_search, "search_openverse_images", fake_openverse)
    monkeypatch.setattr(image_search, "search_commons_images", fake_commons)
    monkeypatch.setattr(image_search, "translate_korean_query_to_english", fake_translate)

    results, searched_query = await search_images_with_query("거북이", 3, "auto")

    assert [candidate.title for candidate in results] == ["Turtle"]
    assert searched_query == "turtle"


def test_openverse_candidates_are_normalized() -> None:
    results = normalize_openverse_images(
        {
            "results": [
                {
                    "id": "abc",
                    "title": "Rabbit photo",
                    "url": "https://example.com/rabbit.jpg",
                    "thumbnail": "https://example.com/rabbit-thumb.jpg",
                    "foreign_landing_url": "https://openverse.org/image/abc",
                    "creator": "Jane",
                    "license": "by",
                    "license_version": "4.0",
                    "license_url": "https://creativecommons.org/licenses/by/4.0/",
                }
            ]
        }
    )

    assert [candidate.model_dump() for candidate in results] == [
        {
            "id": "openverse:abc",
            "title": "Rabbit photo",
            "image_url": "https://example.com/rabbit.jpg",
            "thumbnail_url": "https://example.com/rabbit-thumb.jpg",
            "source_url": "https://openverse.org/image/abc",
            "credit": "Jane",
            "license": "CC BY 4.0",
            "license_url": "https://creativecommons.org/licenses/by/4.0/",
            "provider": "openverse",
        }
    ]


@pytest.mark.anyio
async def test_resolve_image_ignores_unreachable_remote_image(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    class FailingClient:
        def __init__(self, *args: object, **kwargs: object) -> None:
            pass

        async def __aenter__(self) -> "FailingClient":
            return self

        async def __aexit__(self, *args: object) -> None:
            pass

        async def get(self, value: str) -> object:
            raise httpx.ConnectError("network unavailable")

    monkeypatch.setattr(generators.httpx, "AsyncClient", FailingClient)

    assert await generators.resolve_image("https://example.com/rabbit.jpg") is None


def docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def tiny_png_data_uri(width: int = 12, height: int = 12) -> str:
    buffer = BytesIO()
    Image.new("RGB", (width, height), color=(240, 120, 80)).save(buffer, format="PNG")
    encoded = b64encode(buffer.getvalue()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def archive_has_media(content: bytes, prefix: str) -> bool:
    with ZipFile(BytesIO(content)) as archive:
        return any(name.startswith(prefix) for name in archive.namelist())


def media_dimensions(content: bytes, prefix: str) -> set[tuple[int, int]]:
    dimensions: set[tuple[int, int]] = set()
    with ZipFile(BytesIO(content)) as archive:
        for name in archive.namelist():
            if name.startswith(prefix) and name.lower().endswith(IMAGE_MEDIA_EXTENSIONS):
                with Image.open(BytesIO(archive.read(name))) as image:
                    dimensions.add(image.size)
    return dimensions


def docx_picture_widths(content: bytes) -> set[int]:
    document = Document(BytesIO(content))
    return {shape.width for shape in document.inline_shapes}


def pptx_slide_xmls(content: bytes) -> list[str]:
    with ZipFile(BytesIO(content)) as archive:
        slide_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("ppt/slides/slide") and name.endswith(".xml")
        )
        return [archive.read(name).decode("utf-8") for name in slide_names]
